"""Convert client-intake.md to a formatted Word document."""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_run_with_formatting(para, text):
    """Add a run, handling **bold** and `code` inline."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            para.add_run(part)


def parse_inline(text):
    """Strip markdown inline syntax, return plain text for tables."""
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def build_doc(md_path, out_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Styles
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(10.5)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # --- H1
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:], level=1)
            p.runs[0].font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

        # --- H2
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
            p.runs[0].font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

        # --- H3
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)

        # --- horizontal rule
        elif re.match(r"^-{3,}$", line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run()
            run.add_break()

        # --- blockquote
        elif line.startswith("> "):
            p = doc.add_paragraph(style="Quote")
            content = line[2:]
            # strip leading ⚠️ emoji prefix if present
            add_run_with_formatting(p, content)

        # --- table (detect by leading |)
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].rstrip("\n").startswith("|"):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1

            # filter out separator rows (---|---)
            data_rows = [
                r for r in table_lines if not re.match(r"^\|[-| :]+\|$", r)
            ]

            if not data_rows:
                continue

            cols = [c.strip() for c in data_rows[0].split("|") if c.strip()]
            num_cols = len(cols)
            table = doc.add_table(rows=len(data_rows), cols=num_cols)
            table.style = "Table Grid"

            for row_idx, row_text in enumerate(data_rows):
                cells = [c.strip() for c in row_text.split("|")]
                cells = [c for c in cells if c != ""]  # drop empty edge cells
                for col_idx, cell_text in enumerate(cells[:num_cols]):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = parse_inline(cell_text)
                    if row_idx == 0:
                        set_cell_bg(cell, "1A3C5E")
                        run = cell.paragraphs[0].runs
                        if run:
                            run[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run[0].bold = True
                    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

            doc.add_paragraph()  # spacing after table
            continue  # already advanced i in inner loop

        # --- checkbox list item
        elif re.match(r"^- \[[ x]\] ", line):
            p = doc.add_paragraph(style="List Bullet")
            checked = line[3] == "x"
            checkbox = "☑ " if checked else "☐ "
            add_run_with_formatting(p, checkbox + line[6:])

        # --- numbered list item
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            content = re.sub(r"^\d+\. ", "", line)
            add_run_with_formatting(p, content)

        # --- bullet list
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_run_with_formatting(p, line[2:])

        # --- blank line
        elif line.strip() == "":
            pass  # skip, word handles spacing

        # --- italic/plain paragraph
        elif line.startswith("_") and line.endswith("_"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("_"))
            run.italic = True

        # --- bold standalone line (e.g. **Services (tick all...):**)
        elif line.startswith("**") and line.endswith(":**"):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-3] + ":")
            run.bold = True

        # --- plain paragraph
        else:
            p = doc.add_paragraph()
            add_run_with_formatting(p, line)

        i += 1

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_doc(
        "docs/client-intake.md",
        "docs/client-intake.docx",
    )
